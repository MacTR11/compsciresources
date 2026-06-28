#!/usr/bin/env python3
"""
Convert the Markdown resources into editable Word (.docx) files.

Builds clean DOCX directly with python-docx (no LibreOffice/pandoc needed),
handling headings, bold/italic/inline-code, tables, bullet/numbered lists,
fenced code blocks, blockquotes, horizontal rules and <details> answer
sections (rendered open so answers are visible/editable).

Usage:
    python3 revision-tools/build-docx.py                     # build the default set
    python3 revision-tools/build-docx.py <srcdir> <outdir>   # one folder
"""
import sys, os, re, glob

try:
    import docx
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx not installed. Run: pip install python-docx")

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

INLINE = re.compile(r'(\*\*.+?\*\*|__.+?__|\*.+?\*|`.+?`)')


def add_runs(paragraph, text):
    """Add text to a paragraph, parsing **bold**, *italic* and `code`."""
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)   # links -> just the label
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    for part in INLINE.split(text):
        if not part:
            continue
        if (part.startswith('**') and part.endswith('**')) or (part.startswith('__') and part.endswith('__')):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.italic = True
        else:
            paragraph.add_run(part)


def shade(cell_or_para, fill):
    el = cell_or_para._tc if hasattr(cell_or_para, '_tc') else cell_or_para._p
    pr = el.get_or_add_tcPr() if hasattr(cell_or_para, '_tc') else el.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill)
    pr.append(shd)


def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    # force any collapsible answers open, drop the html tags
    text = '\n'.join(lines)
    text = re.sub(r'</?details[^>]*>', '', text)
    text = text.replace('<summary>', '**').replace('</summary>', '**')
    lines = text.split('\n')

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        # fenced code block
        if line.lstrip().startswith('```'):
            i += 1
            code = []
            while i < n and not lines[i].lstrip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            shade(p, 'F2F2F2')
            r = p.add_run('\n'.join(code))
            r.font.name = 'Consolas'; r.font.size = Pt(9)
            continue

        # table (header row followed by |---| separator)
        if line.strip().startswith('|') and i + 1 < n and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i+1]) and '-' in lines[i+1]:
            block = []
            while i < n and lines[i].strip().startswith('|'):
                block.append(lines[i]); i += 1
            def cells(row):
                row = row.strip()
                if row.startswith('|'): row = row[1:]
                if row.endswith('|'): row = row[:-1]
                return [c.strip() for c in row.split('|')]
            header = cells(block[0])
            data = [cells(r) for r in block[2:]]
            tbl = doc.add_table(rows=1, cols=len(header)); tbl.style = 'Table Grid'
            for j, h in enumerate(header):
                c = tbl.rows[0].cells[j]; c.paragraphs[0].text = ''
                add_runs(c.paragraphs[0], h)
                for run in c.paragraphs[0].runs: run.bold = True
                shade(c, 'D9E2F3')
            for drow in data:
                cs = tbl.add_row().cells
                for j in range(len(header)):
                    cs[j].paragraphs[0].text = ''
                    add_runs(cs[j].paragraphs[0], drow[j] if j < len(drow) else '')
            doc.add_paragraph()
            continue

        # headings
        m = re.match(r'^(#{1,6})\s+(.*)$', line)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 4))
            add_runs(h, m.group(2))
            i += 1
            continue

        # horizontal rule
        if re.match(r'^\s*([-*_])\1{2,}\s*$', line):
            p = doc.add_paragraph(); pr = p._p.get_or_add_pPr()
            pb = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'AAAAAA')
            pb.append(bottom); pr.append(pb)
            i += 1
            continue

        # blockquote
        if line.lstrip().startswith('>'):
            content = re.sub(r'^\s*>\s?', '', line)
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            shade(p, 'FFF8E7')
            add_runs(p, content)
            i += 1
            continue

        # bullet list
        m = re.match(r'^(\s*)[-*+]\s+(.*)$', line)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, m.group(2)); i += 1
            continue
        # numbered list
        m = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
        if m:
            p = doc.add_paragraph(style='List Number')
            add_runs(p, m.group(2)); i += 1
            continue

        # blank
        if not line.strip():
            i += 1
            continue

        # normal paragraph
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)


def build_folder(srcdir, outdir):
    mds = sorted(f for f in glob.glob(os.path.join(srcdir, '*.md'))
                 if os.path.basename(f).lower() != 'readme.md')
    made = 0
    for md in mds:
        out = os.path.join(outdir, os.path.splitext(os.path.basename(md))[0] + '.docx')
        convert(md, out)
        made += 1
    return made, [os.path.basename(m) for m in mds]


DEFAULT_JOBS = [
    ('component-01-computer-systems',                 'Word-Documents/01-Computer-Systems'),
    ('component-02-algorithms-and-programming',       'Word-Documents/02-Algorithms-and-Programming'),
    ('component-03-04-programming-project',           'Word-Documents/03-04-Programming-Project'),
    ('revision-tools/guides',                         'Word-Documents/Guides'),
    ('revision-tools/knowledge-organisers',           'Word-Documents/Knowledge-Organisers'),
    ('revision-tools/worksheets',                     'Word-Documents/Worksheets'),
    ('revision-tools/revision-games',                 'Word-Documents/Revision-Games'),
    ('revision-tools/subtopic-quizzes',               'Word-Documents/Subtopic-Quizzes'),
    ('revision-tools/mini-papers',                    'Word-Documents/Mini-Papers'),
    ('revision-tools/mock-papers',                    'Word-Documents/Mock-Papers'),
    ('revision-tools/a-star-pack',                    'Word-Documents/A-Star-Pack'),
    ('revision-tools/programming-workbook',           'Word-Documents/Programming-Workbook'),
    ('revision-tools/practice-questions',             'Word-Documents/Practice-Questions'),
    ('revision-tools/scheme-of-work',                 'Word-Documents/Scheme-of-Work'),
]


def main():
    if len(sys.argv) == 3:
        jobs = [(sys.argv[1], sys.argv[2])]
    else:
        jobs = [(os.path.join(REPO, s), os.path.join(REPO, o)) for s, o in DEFAULT_JOBS]
    total = 0
    for src, out in jobs:
        if not os.path.isdir(src):
            print(f"skip (missing): {src}"); continue
        made, names = build_folder(src, out)
        print(f"{os.path.relpath(src, REPO)} -> {os.path.relpath(out, REPO)}: {made} docx")
        total += made
    print(f"\nDone. {total} DOCX generated.")


if __name__ == '__main__':
    main()
